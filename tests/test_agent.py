"""End-to-end tests. These run fully offline via the deterministic engine."""
from __future__ import annotations

from io import BytesIO

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.agent import AgentOrchestrator
from app.agent.templates import detect_document_type
from app.llm.client import try_parse_json
from app.main import app

client = TestClient(app)


# --------------------------------------------------------------------------- #
# Unit-ish tests
# --------------------------------------------------------------------------- #
def test_detect_document_type():
    assert detect_document_type("write a proposal for a new CRM") == "proposal"
    assert detect_document_type("meeting minutes for the standup") == "meeting_minutes"
    assert detect_document_type("technical design doc for a payments service") == "technical_design"
    assert detect_document_type("an SOP for onboarding") == "sop"
    assert detect_document_type("tell me about widgets") == "business_report"  # default


def test_try_parse_json_handles_fences_and_prose():
    assert try_parse_json('```json\n{"a": 1}\n```') == {"a": 1}
    assert try_parse_json('Sure! Here it is: {"a": 1} hope that helps') == {"a": 1}
    assert try_parse_json("not json at all") is None


# --------------------------------------------------------------------------- #
# Orchestrator: the two required scenarios
# --------------------------------------------------------------------------- #
def test_standard_request_end_to_end():
    orch = AgentOrchestrator()
    result = orch.run("Write a business proposal for a customer support chatbot for a mid-sized retailer.")

    assert result.plan.document_type == "proposal"
    assert len(result.plan.tasks) >= 5
    # Every planned section is drafted with real content.
    assert len(result.contents) == len(result.plan.sections)
    assert all(c.word_count() > 0 for c in result.contents)
    assert result.word_count > 250
    # A valid .docx is produced.
    doc = Document(BytesIO(result.docx_bytes))
    assert len(doc.paragraphs) > 10
    assert any(t.rows for t in doc.tables)  # at least one populated table


def test_complex_ambiguous_request_makes_assumptions():
    orch = AgentOrchestrator()
    result = orch.run(
        "We need something for the leadership offsite next quarter about improving "
        "delivery speed without hurting quality — figure out the format."
    )
    # Agent decided a document type and recorded assumptions for the ambiguity.
    assert result.plan.document_type
    assert len(result.plan.assumptions) >= 1
    assert result.reflection.quality_score >= 0
    assert result.docx_bytes[:2] == b"PK"  # docx is a zip


# --------------------------------------------------------------------------- #
# API tests
# --------------------------------------------------------------------------- #
def test_health():
    r = client.get("/health")
    assert r.status_code == 200
    assert r.json()["status"] == "ok"


def test_agent_endpoint_and_download():
    r = client.post("/agent", json={"request": "Create a project plan for migrating to the cloud."})
    assert r.status_code == 200
    body = r.json()
    assert body["document_type"] == "project_plan"
    assert body["plan"]["tasks"]
    assert body["reflection"]["quality_score"] >= 0
    assert "trace" in body and len(body["trace"]) >= 4

    doc_id = body["document"]["id"]
    dl = client.get(f"/documents/{doc_id}")
    assert dl.status_code == 200
    assert dl.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument"
    )
    Document(BytesIO(dl.content))  # opens without error


@pytest.mark.parametrize("bad", ["", "   ", "x" * 5000])
def test_validation_rejects_bad_requests(bad):
    r = client.post("/agent", json={"request": bad})
    assert r.status_code == 422


def test_base64_inline_option():
    r = client.post("/agent", json={"request": "Draft an SOP for password resets.", "include_base64": True})
    assert r.status_code == 200
    assert r.json()["document"]["base64"]
