"""Render a DocumentPlan + drafted sections into a polished .docx (in memory).

Produces a professional layout: branded title page, metadata table, contents
list, numbered sections with paragraphs / bullets / shaded-header tables, an
appendix recording the agent's assumptions and self-check result, and a footer
with page numbers. Returns raw bytes so the API can stream or persist them.
"""
from __future__ import annotations

from datetime import date
from io import BytesIO
from typing import List, Optional

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from ..schemas import Block, DocumentPlan, ReflectionResult, SectionContent

# Brand palette
_PRIMARY = RGBColor(0x1F, 0x3A, 0x5F)   # deep navy
_ACCENT = RGBColor(0x2E, 0x6D, 0xB4)    # blue
_MUTED = RGBColor(0x6B, 0x72, 0x80)     # grey
_HEADER_FILL = "1F3A5F"
_ZEBRA_FILL = "EEF3F9"


# --------------------------------------------------------------------------- #
# low-level XML helpers
# --------------------------------------------------------------------------- #
def _shade_cell(cell, fill_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _bottom_border(paragraph, color_hex: str = "2E6DB4", size: str = "12") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    borders.append(bottom)
    p_pr.append(borders)


def _add_page_number(paragraph) -> None:
    """Insert a PAGE field into a footer paragraph."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


# --------------------------------------------------------------------------- #
# styling
# --------------------------------------------------------------------------- #
def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for level, size, color in (("Heading 1", 15, _PRIMARY), ("Heading 2", 12, _ACCENT)):
        style = doc.styles[level]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def _footer(doc: Document, plan: DocumentPlan) -> None:
    footer = doc.sections[0].footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(
        f"Autonomous Document Agent  ·  {plan.title[:60]}  ·  Page "
    )
    run.font.size = Pt(8)
    run.font.color.rgb = _MUTED
    _add_page_number(para)


# --------------------------------------------------------------------------- #
# page sections
# --------------------------------------------------------------------------- #
def _title_page(doc: Document, plan: DocumentPlan) -> None:
    for _ in range(4):
        doc.add_paragraph()

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tag.add_run(_doc_type_label(plan).upper())
    tr.font.size = Pt(12)
    tr.font.bold = True
    tr.font.color.rgb = _ACCENT

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    trun = title.add_run(plan.title)
    trun.font.size = Pt(28)
    trun.font.bold = True
    trun.font.color.rgb = _PRIMARY
    _bottom_border(title)

    if plan.subtitle:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        srun = sub.add_run(plan.subtitle)
        srun.font.size = Pt(12)
        srun.italic = True
        srun.font.color.rgb = _MUTED

    for _ in range(2):
        doc.add_paragraph()

    meta = doc.add_table(rows=0, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ("Document Type", _doc_type_label(plan)),
        ("Prepared For", plan.audience),
        ("Date", _format_date(plan.date)),
        ("Prepared By", "Autonomous Document Agent"),
    ]
    for label, value in rows:
        cells = meta.add_row().cells
        run = cells[0].paragraphs[0].add_run(label)
        run.bold = True
        run.font.color.rgb = _PRIMARY
        cells[1].paragraphs[0].add_run(value)

    doc.add_page_break()


def _contents_page(doc: Document, contents: List[SectionContent]) -> None:
    h = doc.add_paragraph()
    hr = h.add_run("Contents")
    hr.font.size = Pt(16)
    hr.font.bold = True
    hr.font.color.rgb = _PRIMARY
    _bottom_border(h)
    doc.add_paragraph()
    for i, content in enumerate(contents, start=1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}.  {content.heading}")
        run.font.size = Pt(11)
    doc.add_page_break()


def _render_block(doc: Document, block: Block) -> None:
    if block.type == "paragraph" and block.text:
        doc.add_paragraph(block.text)
    elif block.type == "subheading" and block.text:
        doc.add_heading(block.text, level=2)
    elif block.type == "bullets" and block.items:
        for item in block.items:
            doc.add_paragraph(item, style="List Bullet")
    elif block.type == "table" and block.columns and block.rows:
        _render_table(doc, block)


def _render_table(doc: Document, block: Block) -> None:
    table = doc.add_table(rows=1, cols=len(block.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, col in enumerate(block.columns):
        _shade_cell(hdr[i], _HEADER_FILL)
        para = hdr[i].paragraphs[0]
        run = para.add_run(str(col))
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    for r, row in enumerate(block.rows):
        cells = table.add_row().cells
        for i in range(len(block.columns)):
            value = str(row[i]) if i < len(row) else ""
            para = cells[i].paragraphs[0]
            run = para.add_run(value)
            run.font.size = Pt(10)
            if r % 2 == 1:
                _shade_cell(cells[i], _ZEBRA_FILL)

    if block.caption:
        cap = doc.add_paragraph()
        crun = cap.add_run(block.caption)
        crun.italic = True
        crun.font.size = Pt(9)
        crun.font.color.rgb = _MUTED
    doc.add_paragraph()


def _sections(doc: Document, contents: List[SectionContent]) -> None:
    for i, content in enumerate(contents, start=1):
        doc.add_heading(f"{i}. {content.heading}", level=1)
        for block in content.blocks:
            _render_block(doc, block)


def _appendix(doc: Document, plan: DocumentPlan, reflection: ReflectionResult) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix A · Assumptions & Quality Check", level=1)

    if plan.assumptions:
        doc.add_heading("Assumptions made by the agent", level=2)
        for a in plan.assumptions:
            doc.add_paragraph(a, style="List Bullet")

    doc.add_heading("Automated self-check", level=2)
    doc.add_paragraph(
        f"Quality score: {reflection.quality_score}/100 · "
        f"Passed: {'Yes' if reflection.passed else 'No'} · "
        f"Revised after review: {'Yes' if reflection.revised else 'No'}."
    )
    if reflection.issues:
        doc.add_paragraph("Issues detected during self-check:")
        for issue in reflection.issues:
            doc.add_paragraph(issue, style="List Bullet")
    else:
        doc.add_paragraph("No structural issues were detected during the self-check.")

    note = doc.add_paragraph()
    nrun = note.add_run(
        "This document was generated autonomously. Figures, names and dates are "
        "illustrative mock data unless supplied in the original request."
    )
    nrun.italic = True
    nrun.font.size = Pt(9)
    nrun.font.color.rgb = _MUTED


# --------------------------------------------------------------------------- #
# helpers
# --------------------------------------------------------------------------- #
def _doc_type_label(plan: DocumentPlan) -> str:
    from ..agent.templates import DOC_TYPE_LABELS

    return DOC_TYPE_LABELS.get(plan.document_type, "Business Document")


def _format_date(iso: Optional[str]) -> str:
    if not iso:
        return date.today().strftime("%d %B %Y")
    try:
        return date.fromisoformat(iso).strftime("%d %B %Y")
    except Exception:  # noqa: BLE001
        return iso


# --------------------------------------------------------------------------- #
# public API
# --------------------------------------------------------------------------- #
def build_docx(plan: DocumentPlan, contents: List[SectionContent], reflection: ReflectionResult) -> bytes:
    doc = Document()
    _configure_styles(doc)
    _footer(doc, plan)
    _title_page(doc, plan)
    _contents_page(doc, contents)
    _sections(doc, contents)
    _appendix(doc, plan, reflection)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
